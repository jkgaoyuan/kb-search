import os
import re
from typing import Tuple
import mistune
from bs4 import BeautifulSoup
import bleach
from docx import Document as DocxDocument
from app.config import get_settings
from app.utils.text_utils import strip_markdown, generate_summary

settings = get_settings()

# 允许的标签和属性
ALLOWED_TAGS = [
    'p', 'br', 'div', 'span', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
    'ul', 'ol', 'li', 'strong', 'em', 'code', 'pre', 'blockquote',
    'a', 'img', 'table', 'thead', 'tbody', 'tr', 'th', 'td'
]
ALLOWED_ATTRIBUTES = {
    'a': ['href', 'title'],
    'img': ['src', 'alt', 'title'],
    'code': ['class'],
    'pre': ['class'],
}
ALLOWED_PROTOCOLS = ['http', 'https', 'mailto']


class DocumentParser:
    """文档解析器：支持 Markdown / HTML / DOCX"""

    def parse(self, file_path: str, file_type: str) -> Tuple[str, str]:
        """
        解析文档，返回 (纯文本, HTML)

        Args:
            file_path: 文件路径
            file_type: markdown | html | docx

        Returns:
            (纯文本内容, HTML内容)
        """
        if file_type == "markdown":
            return self._parse_markdown(file_path)
        elif file_type == "html":
            return self._parse_html(file_path)
        elif file_type == "docx":
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    def _parse_markdown(self, file_path: str) -> Tuple[str, str]:
        with open(file_path, 'r', encoding='utf-8') as f:
            raw_content = f.read()

        # 转换为HTML
        html = mistune.create_markdown()(raw_content)

        # 提取纯文本
        plain_text = strip_markdown(raw_content)

        return plain_text, html

    def _parse_html(self, file_path: str) -> Tuple[str, str]:
        with open(file_path, 'r', encoding='utf-8') as f:
            raw_content = f.read()

        # 清理HTML
        soup = BeautifulSoup(raw_content, 'html.parser')

        # 移除script和style
        for tag in soup(['script', 'style', 'iframe', 'object', 'embed']):
            tag.decompose()

        # 获取body内容
        body = soup.find('body')
        if body:
            html_content = str(body)
        else:
            html_content = str(soup)

        # 清理危险内容
        clean_html = bleach.clean(
            html_content,
            tags=ALLOWED_TAGS,
            attributes=ALLOWED_ATTRIBUTES,
            protocols=ALLOWED_PROTOCOLS,
            strip=True
        )

        # 提取纯文本
        plain_text = soup.get_text(separator='\n', strip=True)

        return plain_text, clean_html

    def _parse_docx(self, file_path: str) -> Tuple[str, str]:
        doc = DocxDocument(file_path)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        plain_text = '\n'.join(paragraphs)

        # DOCX无原生HTML，简单包装为段落
        html_parts = ['<div class="docx-content">']
        for para in doc.paragraphs:
            if para.text.strip():
                html_parts.append(f'<p>{self._escape_html(para.text)}</p>')
        html_parts.append('</div>')
        html = '\n'.join(html_parts)

        return plain_text, html

    @staticmethod
    def _escape_html(text: str) -> str:
        """转义HTML特殊字符"""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;'))

    def detect_file_type(self, filename: str, content_type: str = None) -> str:
        """根据文件名和MIME类型检测文件类型"""
        ext = os.path.splitext(filename.lower())[1]

        type_map = {
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.html': 'html',
            '.htm': 'html',
            '.docx': 'docx',
        }

        if ext in type_map:
            return type_map[ext]

        # 根据MIME类型兜底
        mime_map = {
            'text/markdown': 'markdown',
            'text/html': 'html',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        }

        if content_type and content_type in mime_map:
            return mime_map[content_type]

        raise ValueError(f"无法识别文件类型: {filename} (MIME: {content_type})")


# 全局解析器实例
parser = DocumentParser()
